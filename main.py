import os
import requests
import gradio as gr
from docx import Document
from dotenv import load_dotenv
import random
from datetime import datetime

load_dotenv()

API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
API_URL = "https://api.deepseek.com/v1/chat/completions"
MODEL_NAME = "deepseek-v4-pro"

CATEGORIES = ["教育", "劳动", "校园其他", "社区服务", "公共安全", "环境卫生", "邻里纠纷", "投诉建议", "其他"]

def mock_classify(content):
    keywords = {
        "教育": ["学校", "老师", "学生", "考试", "课程", "作业", "补习班", "学区"],
        "劳动": ["工资", "加班", "老板", "工作", "辞职", "合同", "社保", "加班"],
        "校园其他": ["校园", "宿舍", "食堂", "社团", "运动会"],
        "社区服务": ["物业", "居委会", "维修", "设施", "停车", "门禁"],
        "公共安全": ["安全", "消防", "交通", "事故", "盗窃", "诈骗"],
        "环境卫生": ["垃圾", "污染", "噪音", "卫生", "绿化", "厕所"],
        "邻里纠纷": ["邻居", "噪音", "漏水", "宠物", "纠纷"],
        "投诉建议": ["投诉", "建议", "反映", "问题", "意见"]
    }
    for category, kw_list in keywords.items():
        for kw in kw_list:
            if kw in content:
                return category
    return random.choice(CATEGORIES)

def mock_summarize(content):
    sentences = [s.strip() for s in content.replace('。', '').replace('！', '').replace('？', '').split('，') if s.strip()]
    if len(sentences) <= 3:
        return content[:200] + "..." if len(content) > 200 else content
    summary = "；".join(sentences[:3]) + "。"
    return summary[:200]

def mock_evaluate(content):
    score = random.randint(3, 8)
    reasons = [
        {"影响人数": ["较少（1-10人）", "中等（10-100人）", "较多（100人以上）"],
         "传播潜力": ["较低，内容较私人", "中等，可能在小范围传播", "较高，容易引发关注"],
         "历史参考": ["未见类似事件", "有少量类似案例", "曾有多次类似事件"],
         "潜在影响": ["影响有限", "可能引发局部关注", "可能引发广泛讨论"]},
        {"影响人数": ["中等（10-100人）", "较多（100人以上）", "较多（100人以上）"],
         "传播潜力": ["中等，可能在小范围传播", "较高，容易引发关注", "较高，容易引发关注"],
         "历史参考": ["有少量类似案例", "曾有多次类似事件", "曾有多次类似事件"],
         "潜在影响": ["可能引发局部关注", "可能引发广泛讨论", "可能引发广泛讨论"]}
    ]
    level = 0 if score <= 5 else 1
    reason_set = reasons[level]
    return f"""热度评分：{score}分
评估理由：
- 影响人数：{reason_set["影响人数"][score % 3]}
- 传播潜力：{reason_set["传播潜力"][score % 3]}
- 历史参考：{reason_set["历史参考"][score % 3]}
- 潜在影响：{reason_set["潜在影响"][score % 3]}"""

def mock_title(content):
    keywords = ["事件", "问题", "投诉", "建议", "求助", "反映"]
    title = content[:15].strip()
    if not title:
        title = "社区事件"
    elif not any(k in title for k in keywords):
        title += "事件"
    return title

def call_deepseek(prompt, max_tokens=2048):
    if not API_KEY or API_KEY.strip() == "":
        return "API_KEY_NOT_SET"
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {API_KEY}"
    }
    
    data = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system", "content": "你是一个专业的社区事件分析助手，擅长处理投稿内容的分类、摘要和影响评估。"},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": max_tokens,
        "temperature": 0.3
    }
    
    try:
        response = requests.post(API_URL, headers=headers, json=data, timeout=60)
        
        if response.status_code == 402:
            return "API_ERROR_402"
        elif response.status_code == 401:
            return "API_ERROR_401"
        elif not response.ok:
            return f"API_ERROR_{response.status_code}"
        
        result = response.json()
        return result["choices"][0]["message"]["content"].strip()
    except requests.exceptions.RequestException as e:
        return f"NETWORK_ERROR_{str(e)}"
    except Exception as e:
        return f"UNKNOWN_ERROR_{str(e)}"

def analyze_post(content, use_mock=False):
    original_content = content
    if not content or content.strip() == "":
        return "请输入事件描述内容", "", "", ""
    
    if use_mock:
        category = mock_classify(content)
        summary = mock_summarize(content)
        evaluation = mock_evaluate(content)
        title = mock_title(content)
        return title, category, summary, evaluation
    
    classify_prompt = f"""
请对以下社区投稿内容进行分类，只能从以下类别中选择一个：教育、劳动、校园其他、社区服务、公共安全、环境卫生、邻里纠纷、投诉建议、其他

投稿内容：
{content}

类别：
"""
    category = call_deepseek(classify_prompt, max_tokens=50)
    
    print(f"DEBUG category API response: '{category}'")
    
    if category.startswith("API_ERROR_") or category.startswith("NETWORK_ERROR_") or category.startswith("UNKNOWN_ERROR_") or category == "API_KEY_NOT_SET":
        return f"错误：{category}", f"错误：{category}", f"错误：{category}", f"错误：{category}"
    
    if not category or category.strip() == "":
        category = mock_classify(original_content)
        print(f"DEBUG category fallback to: '{category}'")
    
    summarize_prompt = f"""
    请将以下社区投稿内容浓缩成3-5句话的核心摘要，要求：
    1. 保留关键人物、时间、地点、事件经过
    2. 突出核心诉求和问题
    3. 语言简洁明了
    
    投稿内容：
    {content}
    """
    summary = call_deepseek(summarize_prompt, max_tokens=300)
    
    if summary.startswith("API_ERROR_") or summary.startswith("NETWORK_ERROR_") or summary.startswith("UNKNOWN_ERROR_") or summary == "API_KEY_NOT_SET":
        return f"错误：{summary}", f"错误：{summary}", f"错误：{summary}", f"错误：{summary}"
    
    evaluate_prompt = f"""
    请对以下社区投稿内容进行影响评估：
    1. 给出一个热度分数（1-10分），1分最低，10分最高
    2. 说明评分理由，包括：
       - 影响人数多少
       - 是否容易传播
       - 是否有类似事件先例
       - 潜在社会影响
    
    投稿内容：
    {content}
    
    输出格式：
    热度评分：X分
    评估理由：
    - 影响人数：...
    - 传播潜力：...
    - 历史参考：...
    - 潜在影响：...
    """
    evaluation = call_deepseek(evaluate_prompt, max_tokens=400)
    
    if evaluation.startswith("API_ERROR_") or evaluation.startswith("NETWORK_ERROR_") or evaluation.startswith("UNKNOWN_ERROR_") or evaluation == "API_KEY_NOT_SET":
        return f"错误：{evaluation}", f"错误：{evaluation}", f"错误：{evaluation}", f"错误：{evaluation}"
    
    title_prompt = f"""
请为以下社区投稿内容生成一个简洁明了的事件标题（10-20字），直接输出标题，不要有任何其他内容：

投稿内容：
{content}

标题：
"""
    title = call_deepseek(title_prompt, max_tokens=50)
    
    print(f"DEBUG title API response: '{title}'")
    
    if title.startswith("API_ERROR_") or title.startswith("NETWORK_ERROR_") or title.startswith("UNKNOWN_ERROR_") or title == "API_KEY_NOT_SET":
        return f"错误：{title}", f"错误：{title}", f"错误：{title}", f"错误：{title}"
    
    if not title or title.strip() == "":
        title = original_content[:15].strip() + "事件"
        print(f"DEBUG title fallback to: '{title}'")
    
    return title, category, summary, evaluation

def generate_report(title, category, summary, evaluation):
    report = f"""# 社区投稿分析报告

## 事件标题
{title}

## 事件类型
{category}

## 核心摘要
{summary}

## 影响评估
{evaluation}

---
*报告生成时间：{os.popen('date /t').read().strip()}*
"""
    return report

def export_to_txt(report):
    today = datetime.now()
    date_folder = today.strftime("%Y-%m-%d")
    timestamp = today.strftime("%H-%M-%S")
    
    output_dir = os.path.join("输出报告", date_folder)
    os.makedirs(output_dir, exist_ok=True)
    
    file_name = f"社区投稿分析报告_{timestamp}.txt"
    file_path = os.path.join(output_dir, file_name)
    
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(report)
    
    return file_path

def export_to_word(report):
    from io import BytesIO
    
    today = datetime.now()
    date_folder = today.strftime("%Y-%m-%d")
    timestamp = today.strftime("%H-%M-%S")
    
    output_dir = os.path.join("输出报告", date_folder)
    os.makedirs(output_dir, exist_ok=True)
    
    doc = Document()
    doc.add_heading('社区投稿分析报告', 0)
    
    lines = report.split('\n')
    for line in lines:
        if line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('- '):
            doc.add_paragraph(line, style='List Bullet')
        elif line.startswith('# ') or line.startswith('---'):
            pass
        elif line.strip():
            doc.add_paragraph(line)
    
    file_name = f"社区投稿分析报告_{timestamp}.docx"
    file_path = os.path.join(output_dir, file_name)
    
    doc.save(file_path)
    
    return file_path

current_report = ""

def process_input(content, use_mock):
    global current_report
    try:
        global API_KEY
        title, category, summary, evaluation = analyze_post(content, use_mock)
        
        print(f"DEBUG process_input:")
        print(f"  title = '{title}'")
        print(f"  category = '{category}'")
        print(f"  summary = '{summary[:50]}...'")
        print(f"  evaluation = '{evaluation[:50]}...'")
        
        report = generate_report(title, category, summary, evaluation)
        current_report = report
        
        export_to_txt(report)
        export_to_word(report)
        
        return title, category, summary, evaluation, report
    except Exception as e:
        error_msg = f"处理错误: {str(e)}"
        print(f"ERROR - {error_msg}")
        return error_msg, error_msg, error_msg, error_msg, error_msg

def download_txt():
    global current_report
    if not current_report:
        return None
    file_path = export_to_txt(current_report)
    return file_path

def download_word():
    global current_report
    if not current_report:
        return None
    file_path = export_to_word(current_report)
    return file_path

def update_api_key(new_key):
    global API_KEY
    API_KEY = new_key
    return "API密钥已更新"

with gr.Blocks(title="社区投稿AI处理助手") as demo:
    gr.Markdown("# 🏘️ 社区投稿AI处理助手")
    gr.Markdown("自动分类、摘要和影响评估，帮助社区更快响应投稿")
    
    with gr.Row():
        with gr.Column(scale=1):
            with gr.Accordion("⚙️ API配置", open=False):
                api_key_input = gr.Textbox(
                    label="DeepSeek API Key",
                    placeholder="请输入您的DeepSeek API Key...",
                    value=API_KEY,
                    type="password"
                )
                update_key_btn = gr.Button("更新API密钥")
                key_status = gr.Textbox(label="状态", interactive=False)
                
                update_key_btn.click(update_api_key, inputs=[api_key_input], outputs=[key_status])
            
            gr.Markdown("### 📝 事件描述")
            input_text = gr.Textbox(
                label="投稿内容",
                placeholder="请粘贴投稿内容（支持微信/小红书/抖音内容）...",
                lines=12,
                max_lines=15
            )
            
            use_mock_checkbox = gr.Checkbox(
                label="使用模拟模式（无需API密钥）",
                value=False
            )
            
            submit_btn = gr.Button("开始分析", variant="primary", size="lg")
        
        with gr.Column(scale=1):
            with gr.Tabs():
                with gr.Tab("分析结果"):
                    title_output = gr.Textbox(label="事件标题", interactive=False)
                    category_output = gr.Textbox(label="事件类型", interactive=False)
                    summary_output = gr.Textbox(label="核心摘要", lines=5, interactive=False)
                    evaluation_output = gr.Textbox(label="影响评估", lines=8, interactive=False)
                
                with gr.Tab("完整报告"):
                    report_output = gr.Markdown(label="Markdown报告")
    
    with gr.Row():
        with gr.Column(scale=1):
            txt_download_btn = gr.Button("下载TXT报告")
            txt_download = gr.File(label="TXT文件")
        with gr.Column(scale=1):
            word_download_btn = gr.Button("下载Word报告")
            word_download = gr.File(label="Word文件")
    
    with gr.Row():
        gr.Markdown("""
        **💡 使用提示：**
        - 若使用真实API，需确保DeepSeek账户有足够余额
        - 402错误表示账户余额不足，请充值或使用模拟模式
        - 模拟模式可体验完整功能，但结果为随机生成
        """)
    
    submit_btn.click(
        fn=process_input,
        inputs=[input_text, use_mock_checkbox],
        outputs=[title_output, category_output, summary_output, evaluation_output, report_output]
    )
    
    txt_download_btn.click(
        fn=download_txt,
        outputs=[txt_download]
    )
    
    word_download_btn.click(
        fn=download_word,
        outputs=[word_download]
    )

if __name__ == "__main__":
    demo.launch(share=False)
